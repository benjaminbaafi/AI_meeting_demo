"""
Export utilities for generating PDF, DOCX, and JSON exports.
Handles formatting of transcriptions, summaries, and action items.
"""
import json
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, List
import io

from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_CENTER

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


class ExportUtils:
    """Utilities for exporting data in various formats."""
    
    @staticmethod
    def export_transcription_json(transcription_data: Dict[str, Any]) -> bytes:
        """
        Export transcription as JSON.
        
        Args:
            transcription_data: Transcription data
            
        Returns:
            JSON bytes
        """
        return json.dumps(transcription_data, indent=2, default=str).encode('utf-8')
    
    @staticmethod
    def export_transcription_pdf(transcription_data: Dict[str, Any]) -> bytes:
        """
        Export transcription as PDF.
        
        Args:
            transcription_data: Transcription data
            
        Returns:
            PDF bytes
        """
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter,
                              topMargin=0.75*inch, bottomMargin=0.75*inch)
        
        # Styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1a73e8'),
            spaceAfter=30,
            alignment=TA_CENTER
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor('#5f6368'),
            spaceBefore=12,
            spaceAfter=6
        )
        
        # Build content
        content = []
        
        # Title
        content.append(Paragraph("Meeting Transcription", title_style))
        content.append(Spacer(1, 0.3*inch))
        
        # Metadata
        metadata_data = [
            ['Job ID:', transcription_data.get('job_id', 'N/A')],
            ['Duration:', f"{transcription_data.get('duration_seconds', 0):.0f} seconds"],
            ['Word Count:', str(transcription_data.get('word_count', 0))],
            ['Language:', transcription_data.get('language_detected', 'N/A')],
            ['Generated:', datetime.now().strftime('%Y-%m-%d %H:%M:%S')],
        ]
        
        metadata_table = Table(metadata_data, colWidths=[2*inch, 4*inch])
        metadata_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f1f3f4')),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (0, -1), 'RIGHT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('LEFTPADDING', (0, 0), (-1, -1), 12),
            ('RIGHTPADDING', (0, 0), (-1, -1), 12),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ]))
        
        content.append(metadata_table)
        content.append(Spacer(1, 0.4*inch))
        
        # Speakers
        if transcription_data.get('speakers'):
            content.append(Paragraph("Speakers", heading_style))
            speakers_text = ", ".join([s.get('name', 'Unknown') for s in transcription_data['speakers']])
            content.append(Paragraph(speakers_text, styles['Normal']))
            content.append(Spacer(1, 0.2*inch))
        
        # Transcript
        content.append(Paragraph("Transcript", heading_style))
        content.append(Spacer(1, 0.1*inch))
        
        # Formatted text
        if transcription_data.get('formatted_text'):
            lines = transcription_data['formatted_text'].split('\n')
            for line in lines:
                if line.strip():
                    content.append(Paragraph(line, styles['Normal']))
                    content.append(Spacer(1, 0.1*inch))
        
        # Build PDF
        doc.build(content)
        buffer.seek(0)
        return buffer.getvalue()
    
    @staticmethod
    def export_transcription_docx(transcription_data: Dict[str, Any]) -> bytes:
        """
        Export transcription as DOCX.
        
        Args:
            transcription_data: Transcription data
            
        Returns:
            DOCX bytes
        """
        doc = Document()
        
        # Title
        title = doc.add_heading('Meeting Transcription', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata table
        doc.add_heading('Metadata', level=2)
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Light Grid Accent 1'
        
        metadata = [
            ('Job ID', transcription_data.get('job_id', 'N/A')),
            ('Duration', f"{transcription_data.get('duration_seconds', 0):.0f} seconds"),
            ('Word Count', str(transcription_data.get('word_count', 0))),
            ('Language', transcription_data.get('language_detected', 'N/A')),
            ('Generated', datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
        ]
        
        for idx, (label, value) in enumerate(metadata):
            table.rows[idx].cells[0].text = label
            table.rows[idx].cells[1].text = value
        
        doc.add_paragraph()
        
        # Speakers
        if transcription_data.get('speakers'):
            doc.add_heading('Speakers', level=2)
            speakers_text = ", ".join([s.get('name', 'Unknown') for s in transcription_data['speakers']])
            doc.add_paragraph(speakers_text)
        
        # Transcript
        doc.add_heading('Transcript', level=2)
        
        if transcription_data.get('formatted_text'):
            lines = transcription_data['formatted_text'].split('\n')
            for line in lines:
                if line.strip():
                    p = doc.add_paragraph(line)
                    p.paragraph_format.line_spacing = 1.15
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    @staticmethod
    def export_summary_json(summary_data: Dict[str, Any]) -> bytes:
        """Export summary as JSON."""
        return json.dumps(summary_data, indent=2, default=str).encode('utf-8')
    
    @staticmethod
    def export_summary_pdf(summary_data: Dict[str, Any]) -> bytes:
        """
        Export summary as PDF.
        
        Args:
            summary_data: Summary data
            
        Returns:
            PDF bytes
        """
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter,
                              topMargin=0.75*inch, bottomMargin=0.75*inch)
        
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1a73e8'),
            spaceAfter=30,
            alignment=TA_CENTER
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor('#5f6368'),
            spaceBefore=12,
            spaceAfter=6
        )
        
        content = []
        
        # Title
        content.append(Paragraph("Meeting Summary", title_style))
        content.append(Spacer(1, 0.3*inch))
        
        # Summary Type
        summary_type = summary_data.get('summary_type', 'N/A')
        content.append(Paragraph(f"Summary Type: {summary_type}", styles['Normal']))
        content.append(Spacer(1, 0.2*inch))
        
        # Executive Summary
        if summary_data.get('executive_summary'):
            content.append(Paragraph("Executive Summary", heading_style))
            content.append(Paragraph(summary_data['executive_summary'], styles['Normal']))
            content.append(Spacer(1, 0.2*inch))
        
        # Detailed Summary
        if summary_data.get('detailed_summary'):
            content.append(Paragraph("Detailed Summary", heading_style))
            content.append(Paragraph(summary_data['detailed_summary'], styles['Normal']))
            content.append(Spacer(1, 0.2*inch))
        
        # Main Topics
        if summary_data.get('main_topics'):
            content.append(Paragraph("Main Topics", heading_style))
            for topic in summary_data['main_topics']:
                content.append(Paragraph(f"• {topic}", styles['Normal']))
            content.append(Spacer(1, 0.2*inch))
        
        # Action Items
        if summary_data.get('action_items'):
            content.append(PageBreak())
            content.append(Paragraph("Action Items", heading_style))
            
            action_data = [['Description', 'Assignee', 'Deadline', 'Priority']]
            for item in summary_data['action_items']:
                action_data.append([
                    item.get('description', ''),
                    item.get('assignee', 'N/A'),
                    item.get('deadline', 'N/A'),
                    item.get('priority', 'medium')
                ])
            
            action_table = Table(action_data, colWidths=[3*inch, 1.5*inch, 1.2*inch, 0.8*inch])
            action_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1a73e8')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 11),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('LEFTPADDING', (0, 0), (-1, -1), 8),
                ('RIGHTPADDING', (0, 0), (-1, -1), 8),
            ]))
            
            content.append(action_table)
        
        # Key Decisions
        if summary_data.get('key_decisions'):
            content.append(Spacer(1, 0.3*inch))
            content.append(Paragraph("Key Decisions", heading_style))
            for decision in summary_data['key_decisions']:
                content.append(Paragraph(f"• {decision.get('decision', '')}", styles['Normal']))
        
        doc.build(content)
        buffer.seek(0)
        return buffer.getvalue()
    
    @staticmethod
    def export_summary_docx(summary_data: Dict[str, Any]) -> bytes:
        """Export summary as DOCX."""
        doc = Document()
        
        # Title
        title = doc.add_heading('Meeting Summary', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Summary Type
        doc.add_paragraph(f"Summary Type: {summary_data.get('summary_type', 'N/A')}")
        
        # Executive Summary
        if summary_data.get('executive_summary'):
            doc.add_heading('Executive Summary', level=2)
            doc.add_paragraph(summary_data['executive_summary'])
        
        # Detailed Summary
        if summary_data.get('detailed_summary'):
            doc.add_heading('Detailed Summary', level=2)
            doc.add_paragraph(summary_data['detailed_summary'])
        
        # Main Topics
        if summary_data.get('main_topics'):
            doc.add_heading('Main Topics', level=2)
            for topic in summary_data['main_topics']:
                doc.add_paragraph(f"• {topic}")
        
        # Action Items
        if summary_data.get('action_items'):
            doc.add_page_break()
            doc.add_heading('Action Items', level=2)
            
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Light Grid Accent 1'
            
            # Headers
            headers = table.rows[0].cells
            headers[0].text = 'Description'
            headers[1].text = 'Assignee'
            headers[2].text = 'Deadline'
            headers[3].text = 'Priority'
            
            # Data
            for item in summary_data['action_items']:
                row = table.add_row().cells
                row[0].text = item.get('description', '')
                row[1].text = item.get('assignee', 'N/A')
                row[2].text = item.get('deadline', 'N/A')
                row[3].text = item.get('priority', 'medium')
        
        # Key Decisions
        if summary_data.get('key_decisions'):
            doc.add_heading('Key Decisions', level=2)
            for decision in summary_data['key_decisions']:
                doc.add_paragraph(f"• {decision.get('decision', '')}")
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()


def get_export_utils() -> ExportUtils:
    """Get export utils instance."""
    return ExportUtils()
